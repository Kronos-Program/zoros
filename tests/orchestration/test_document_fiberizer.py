from datetime import datetime
from pathlib import Path
from uuid import uuid4

import pytest

pytest.importorskip("docx")
pytest.importorskip("fitz")

from source.core.models.fiber import Fiber
from source.core.models.fibrizer_options import FibrizerOptions
from source.orchestration.fibrizers.document_fiberizer import DocumentFiberizer


def _make_docx(path: Path) -> None:
    import docx

    doc = docx.Document()
    doc.add_heading("Title", level=1)
    doc.add_paragraph("First paragraph.")
    doc.add_paragraph("Second paragraph.")
    doc.save(path)


def _make_pdf(path: Path) -> None:
    import fitz

    doc = fitz.open()
    page1 = doc.new_page()
    page1.insert_text((72, 72), "Page one")
    page2 = doc.new_page()
    page2.insert_text((72, 72), "Page two")
    doc.save(path)
    doc.close()


def _fiber(path: Path) -> Fiber:
    return Fiber(
        id=uuid4(),
        content=str(path),
        type="text",
        metadata={},
        revision_count=0,
        created_at=datetime.utcnow(),
        source="unit",
    )


def test_docx_ingest(tmp_path: Path) -> None:
    docx_path = tmp_path / "sample.docx"
    _make_docx(docx_path)

    df = DocumentFiberizer(FibrizerOptions())
    out = df.fibrize(_fiber(docx_path))
    assert len(out) == 3
    assert all(f.metadata["fold_level"] == 0 for f in out)


def test_pdf_ingest(tmp_path: Path) -> None:
    pdf_path = tmp_path / "sample.pdf"
    _make_pdf(pdf_path)

    df = DocumentFiberizer(FibrizerOptions())
    out = df.fibrize(_fiber(pdf_path))
    assert len(out) == 2
    assert all(f.metadata["fold_level"] == 0 for f in out)
